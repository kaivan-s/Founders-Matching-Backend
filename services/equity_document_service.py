"""
Equity Document Generation Service

Generates jurisdiction-aware Founders' Agreement template documents in PDF and DOCX formats.

The output is a structured starting-point template (NOT a final legal contract) intended
to facilitate discussion between co-founders and serve as a basis for legal review.
Supports US, India, UK, and other jurisdictions with adapted terminology and provisions.

Uses python-docx for DOCX and reportlab for PDF generation.
"""

import io
import uuid
from datetime import datetime, timezone
from typing import Dict, Any, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from config.database import get_supabase, get_supabase_admin
from services.workspace_service import _verify_workspace_access, _get_founder_id, _log_audit
from services.equity_questionnaire_service import get_questionnaire_responses
from utils.logger import log_error, log_info


# ============================================================================
# Jurisdiction-Aware Configuration
# ============================================================================

JURISDICTION_CONFIG = {
    'india': {
        'label': 'India',
        'currency_symbol': '₹',
        'currency_code': 'INR',
        'company_form': 'Private Limited Company (Pvt. Ltd.)',
        'governing_law': 'the laws of India',
        'companies_act': 'Companies Act, 2013',
        'arbitration_act': 'Arbitration and Conciliation Act, 1996',
        'arbitration_seat': 'Bengaluru, India',
        'court_jurisdiction': 'the courts at Bengaluru, India',
        'non_compete_note': (
            'Note: Post-employment non-compete restrictions are generally **not enforceable** under '
            'Section 27 of the Indian Contract Act, 1872. Founders should rely primarily on '
            'non-solicitation and confidentiality clauses for protection.'
        ),
        'tax_note': (
            '- Founders should consider tax implications under the Income Tax Act, 1961, '
            'including capital gains on share transfers and ESOP perquisite tax.\n'
            '- If foreign capital is involved, FEMA and RBI compliance is required.\n'
            '- GST registration may be required once turnover thresholds are met.'
        ),
        'specific_recommendations': [
            'Execute a separate IP Assignment Agreement on stamp paper of appropriate value.',
            'Stamp this document on stamp paper as per the Stamp Act of your state (e.g., ₹100 in Karnataka).',
            'File necessary forms with the Registrar of Companies (RoC) for any equity changes.',
            'Consider executing a Shareholders Agreement (SHA) alongside this Founders Agreement.',
        ],
    },
    'us': {
        'label': 'United States',
        'currency_symbol': '$',
        'currency_code': 'USD',
        'company_form': 'Delaware C-Corporation (recommended for venture-backed startups)',
        'governing_law': 'the laws of the State of Delaware',
        'companies_act': 'Delaware General Corporation Law (DGCL)',
        'arbitration_act': 'Federal Arbitration Act (FAA)',
        'arbitration_seat': 'Wilmington, Delaware',
        'court_jurisdiction': 'the Court of Chancery of the State of Delaware',
        'non_compete_note': (
            'Note: Post-employment non-compete clauses are **not enforceable** in California, '
            'and have limited enforceability in many other states (e.g., North Dakota, Oklahoma, '
            'Massachusetts has restrictions). Founders are strongly advised to rely on '
            'non-solicitation, confidentiality, and trade secret protections.'
        ),
        'tax_note': (
            '- **CRITICAL:** Founders should file an **83(b) election** with the IRS within '
            '**30 days** of receiving restricted stock subject to vesting. Failure to do so can '
            'result in significant adverse tax consequences.\n'
            '- Consider Section 1202 (QSBS) implications for potential capital gains exclusion.\n'
            '- Consult a tax advisor regarding state tax implications of stock grants.'
        ),
        'specific_recommendations': [
            'File 83(b) election with the IRS within 30 days of stock issuance (CRITICAL).',
            'Execute a separate IP Assignment Agreement (PIIA) at incorporation.',
            'Obtain a 409A valuation before issuing options or granting equity to non-founders.',
            'Maintain proper corporate governance: board minutes, stockholder consents, etc.',
        ],
    },
    'uk': {
        'label': 'United Kingdom',
        'currency_symbol': '£',
        'currency_code': 'GBP',
        'company_form': 'Private Limited Company (Ltd.)',
        'governing_law': 'the laws of England and Wales',
        'companies_act': 'Companies Act 2006',
        'arbitration_act': 'Arbitration Act 1996',
        'arbitration_seat': 'London, United Kingdom',
        'court_jurisdiction': 'the courts of England and Wales',
        'non_compete_note': (
            'Note: Under English law, restrictive covenants must be **reasonable in scope, '
            'duration, and geography** to be enforceable. Overly broad non-compete clauses '
            'are routinely struck down. Founders are advised to keep restrictions narrow '
            'and consult counsel on enforceability.'
        ),
        'tax_note': (
            '- Consider EMI (Enterprise Management Incentive) scheme for tax-efficient equity grants.\n'
            '- Be aware of HMRC employment-related securities reporting obligations.\n'
            '- SEIS/EIS qualification may be relevant for early-stage fundraising.'
        ),
        'specific_recommendations': [
            'File necessary forms with Companies House for any share issuance or transfer.',
            'Consider executing a Shareholders Agreement alongside this Founders Agreement.',
            'Obtain HMRC valuation for tax purposes if granting equity at less than market value.',
            'Ensure compliance with the Companies Act 2006 statutory requirements.',
        ],
    },
    'other': {
        'label': 'Other / International',
        'currency_symbol': '$',
        'currency_code': 'USD',
        'company_form': '[Specify your company form, e.g., LLC, Pvt Ltd, GmbH, etc.]',
        'governing_law': '[the laws of your jurisdiction — to be specified]',
        'companies_act': '[Applicable company law — to be specified]',
        'arbitration_act': '[Applicable arbitration law — to be specified]',
        'arbitration_seat': '[City, Country to be specified]',
        'court_jurisdiction': '[Courts of competent jurisdiction in your location]',
        'non_compete_note': (
            'Note: Enforceability of non-compete and restrictive covenants varies '
            'significantly by jurisdiction. Founders are strongly advised to consult '
            'local counsel on what restrictions are valid and reasonable in their location.'
        ),
        'tax_note': (
            '- Founders should consult local tax counsel on the tax implications of equity grants, '
            'vesting, and share transfers in their specific jurisdiction.'
        ),
        'specific_recommendations': [
            'Consult local counsel to adapt this template to your jurisdiction.',
            'Ensure compliance with local company law and registration requirements.',
            'Execute a separate IP Assignment Agreement under local law.',
            'Address local stamp duty, registration, and notarization requirements.',
        ],
    },
}


def _get_jurisdiction_config(jurisdiction: str) -> Dict[str, Any]:
    """Get jurisdiction-specific configuration, defaulting to 'other'."""
    return JURISDICTION_CONFIG.get((jurisdiction or 'other').lower(), JURISDICTION_CONFIG['other'])


def _format_currency(amount: float, jurisdiction: str) -> str:
    """Format a currency amount with the appropriate symbol for the jurisdiction."""
    config = _get_jurisdiction_config(jurisdiction)
    return f"{config['currency_symbol']}{amount:,.0f}"


def _format_jurisdiction_recommendations(jurisdiction: str) -> str:
    """Build the jurisdiction-specific recommendations list."""
    config = _get_jurisdiction_config(jurisdiction)
    items = '\n'.join(f"- {rec}" for rec in config['specific_recommendations'])
    return items


def _format_vesting_description(vesting_terms: Dict[str, Any]) -> str:
    """Format vesting terms into readable text."""
    if not vesting_terms or not vesting_terms.get('has_vesting', True):
        return "The Founders have elected not to implement a vesting schedule. All equity is fully vested upon signing."
    
    years = vesting_terms.get('years', 4)
    cliff_months = vesting_terms.get('cliff_months', 12)
    acceleration = vesting_terms.get('acceleration', 'none')
    
    description = f"Equity shall vest over a period of **{years} years** with a **{cliff_months}-month cliff**.\n\n"
    description += "This means:\n"
    description += f"- No equity vests during the first {cliff_months} months (cliff period)\n"
    description += f"- After the cliff, {100 / years:.1f}% vests on the cliff date\n"
    description += f"- Remaining equity vests monthly over the following {years * 12 - cliff_months} months\n"
    
    if acceleration == 'single_trigger':
        description += "\n**Acceleration:** Single-trigger acceleration applies. Upon a change of control event (acquisition), all unvested equity immediately vests."
    elif acceleration == 'double_trigger':
        description += "\n**Acceleration:** Double-trigger acceleration applies. Unvested equity accelerates only if both (1) a change of control occurs AND (2) the Founder is terminated without cause within 12 months."
    
    return description


def _format_capital_details(
    founder_a_name: str,
    founder_b_name: str,
    founder_a_capital: float,
    founder_b_capital: float
) -> str:
    """Format capital contribution details."""
    total = founder_a_capital + founder_b_capital
    
    if total == 0:
        return "No initial capital contributions have been made by either Founder at the time of this agreement."
    
    details = "The following capital contributions have been made:\n\n"
    details += f"| Founder | Contribution |\n"
    details += f"|---------|-------------|\n"
    
    if founder_a_capital > 0:
        details += f"| {founder_a_name} | ${founder_a_capital:,.0f} |\n"
    else:
        details += f"| {founder_a_name} | $0 |\n"
    
    if founder_b_capital > 0:
        details += f"| {founder_b_name} | ${founder_b_capital:,.0f} |\n"
    else:
        details += f"| {founder_b_name} | $0 |\n"
    
    details += f"| **Total** | **${total:,.0f}** |\n"
    
    return details


def _format_ip_statement(startup_context: Dict[str, Any], founder_a_name: str, founder_b_name: str) -> str:
    """Format IP ownership statement."""
    has_ip = startup_context.get('has_ip', False)
    ip_owner = startup_context.get('ip_owner', '')
    
    if not has_ip:
        return "No pre-existing intellectual property is being contributed to the company. All IP developed after the formation of this partnership shall be owned by the company."
    
    statement = "The following intellectual property is being contributed to the company:\n\n"
    
    if ip_owner == 'founder_a':
        statement += f"- {founder_a_name} contributes pre-existing IP to the company\n"
    elif ip_owner == 'founder_b':
        statement += f"- {founder_b_name} contributes pre-existing IP to the company\n"
    elif ip_owner == 'joint':
        statement += "- Both Founders jointly contribute pre-existing IP to the company\n"
    
    statement += "\nAll contributed IP shall become the property of the company upon signing this agreement."
    
    return statement


def _format_calculation_breakdown(breakdown: Dict[str, Any], founder_a_name: str, founder_b_name: str) -> str:
    """Format calculation breakdown as a markdown table."""
    if not breakdown:
        return "*No calculation breakdown available (equal or custom split selected)*"
    
    weights = breakdown.get('weights', {})
    founder_a = breakdown.get('founder_a', {})
    founder_b = breakdown.get('founder_b', {})
    
    table = "| Factor | Weight | " + founder_a_name + " | " + founder_b_name + " |\n"
    table += "|--------|--------|"
    table += "-" * (len(founder_a_name) + 2) + "|"
    table += "-" * (len(founder_b_name) + 2) + "|\n"
    
    factors = [
        ('time', 'time_commitment', 'Time Commitment'),
        ('capital', 'capital_contribution', 'Capital Contribution'),
        ('expertise', 'domain_expertise', 'Domain Expertise'),
        ('risk', 'risk_taken', 'Risk Taken'),
        ('network', 'network', 'Network'),
        ('idea', 'idea_origination', 'Idea Origination'),
    ]
    
    for score_key, weight_key, label in factors:
        weight_pct = weights.get(weight_key, 0) * 100
        a_score = founder_a.get(score_key, '-')
        b_score = founder_b.get(score_key, '-')
        table += f"| {label} | {weight_pct:.0f}% | {a_score}/10 | {b_score}/10 |\n"
    
    # Add weighted totals
    a_total = founder_a.get('weighted_total', 0)
    b_total = founder_b.get('weighted_total', 0)
    table += f"| **Weighted Total** | - | **{a_total:.2f}** | **{b_total:.2f}** |\n"
    
    return table


def generate_mmd_document(
    scenario: Dict[str, Any],
    questionnaire_responses: Dict[str, Any],
    workspace_title: str
) -> str:
    """
    Generate comprehensive Founders' Agreement template document in MMD (Markdown) format.

    Generates a jurisdiction-aware template (US, India, UK, or Other) with prominent
    "DRAFT — FOR DISCUSSION & LEGAL REVIEW" disclaimers, jurisdiction-specific
    provisions, and a "Before You Sign" checklist.
    
    Args:
        scenario: The approved equity scenario
        questionnaire_responses: Questionnaire responses for both founders
        workspace_title: Workspace/company name
    
    Returns:
        The document content as a string
    """
    # Extract founder info
    founder_a = scenario.get('founder_a', {})
    founder_b = scenario.get('founder_b', {})
    founder_a_name = founder_a.get('name', 'Founder A')
    founder_b_name = founder_b.get('name', 'Founder B')
    
    # Extract percentages - ensure they are correctly extracted from the scenario
    founder_a_percent = float(scenario.get('founder_a_percent', 50))
    founder_b_percent = float(scenario.get('founder_b_percent', 50))
    advisor_percent = float(scenario.get('advisor_percent', 0))
    
    # Validate percentages sum to 100% (with small tolerance for floating point)
    total_percent = founder_a_percent + founder_b_percent + advisor_percent
    if abs(total_percent - 100.0) > 0.01:
        log_error(f"Warning: Equity percentages don't sum to 100%: {founder_a_percent}% + {founder_b_percent}% + {advisor_percent}% = {total_percent}%")
        # Normalize to ensure they sum to 100% (proportionally)
        if total_percent > 0:
            ratio = 100.0 / total_percent
            founder_a_percent = founder_a_percent * ratio
            founder_b_percent = founder_b_percent * ratio
            advisor_percent = advisor_percent * ratio
        else:
            # Fallback to 50/50 if both are 0 or invalid
            founder_a_percent = 50.0
            founder_b_percent = 50.0
            advisor_percent = 0
    
    # Extract vesting terms
    vesting_terms = scenario.get('vesting_terms', {})
    
    # Extract questionnaire data
    founder_a_resp = questionnaire_responses.get('founder_a', {})
    founder_b_resp = questionnaire_responses.get('founder_b', {})
    startup_context = questionnaire_responses.get('startup_context', {})
    
    founder_a_data = founder_a_resp.get('responses', {}) if founder_a_resp else {}
    founder_b_data = founder_b_resp.get('responses', {}) if founder_b_resp else {}
    
    # Extract roles
    founder_a_role_data = founder_a_data.get('role', {})
    founder_b_role_data = founder_b_data.get('role', {})
    founder_a_role = founder_a_role_data.get('title', 'Co-Founder')
    founder_b_role = founder_b_role_data.get('title', 'Co-Founder')
    
    # Extract responsibilities
    founder_a_responsibilities = founder_a_data.get('responsibilities', 'To be determined')
    founder_b_responsibilities = founder_b_data.get('responsibilities', 'To be determined')
    
    # Extract time commitment
    founder_a_time = founder_a_data.get('time_commitment', '')
    founder_b_time = founder_b_data.get('time_commitment', '')
    
    # Extract capital
    founder_a_capital_data = founder_a_data.get('capital_contribution', {})
    founder_b_capital_data = founder_b_data.get('capital_contribution', {})
    founder_a_capital = float(founder_a_capital_data.get('exact_amount', 0) or 0)
    founder_b_capital = float(founder_b_capital_data.get('exact_amount', 0) or 0)
    
    # Extract expertise
    founder_a_expertise = founder_a_data.get('expertise', {})
    founder_b_expertise = founder_b_data.get('expertise', {})
    
    # Extract startup context
    business_stage = startup_context.get('stage', '')
    business_description = startup_context.get('business_description', '')
    
    # Extract calculation breakdown
    breakdown = scenario.get('calculation_breakdown', {})
    
    # Format vesting details
    vesting_years = vesting_terms.get('years', 4)
    cliff_months = vesting_terms.get('cliff_months', 12)
    has_vesting = vesting_terms.get('has_vesting', True)
    acceleration = vesting_terms.get('acceleration', 'none')
    
    # Extract advisor vesting (if advisor equity is set)
    advisor_vesting_years = vesting_terms.get('advisor_vesting_years', 2)
    advisor_cliff_months = vesting_terms.get('advisor_cliff_months', 3)
    
    # Get advisor name from workspace participants (passed in scenario or default)
    advisor_name = scenario.get('advisor_name', 'Project Advisor')
    
    # Format date
    effective_date = datetime.now(timezone.utc)
    date_str = effective_date.strftime('%B %d, %Y')
    day_str = effective_date.strftime('%d')
    month_str = effective_date.strftime('%B')
    year_str = effective_date.strftime('%Y')
    
    # Format responsibilities as list
    def format_responsibilities(resp_text):
        if not resp_text or resp_text == 'To be determined':
            return '- To be determined'
        # Split by newlines or bullets and format
        lines = resp_text.replace('\n', '|').split('|')
        formatted = []
        for line in lines:
            line = line.strip()
            if line:
                if not line.startswith('-'):
                    formatted.append(f"- {line}")
                else:
                    formatted.append(line)
        return '\n'.join(formatted) if formatted else '- To be determined'
    
    founder_a_resp_formatted = format_responsibilities(founder_a_responsibilities)
    founder_b_resp_formatted = format_responsibilities(founder_b_responsibilities)
    
    # Format time commitment
    def format_time_commitment(time_str):
        if not time_str:
            return '[Time commitment to be determined]'
        # Handle various time commitment formats
        time_str_lower = time_str.lower().strip()
        time_map = {
            'full_time': 'Full-time (40+ hours/week)',
            'full_time_now': 'Full-time (40+ hours/week) - Currently active',
            'full_time_soon': 'Full-time (40+ hours/week) - Starting soon',
            'part_time': 'Part-time (20-30 hours/week)',
            'part_time_now': 'Part-time (20-30 hours/week) - Currently active',
            'part_time_soon': 'Part-time (20-30 hours/week) - Starting soon',
            'advisory': 'Advisory (10 hours/week)',
            'advisory_now': 'Advisory (10 hours/week) - Currently active',
            'advisory_soon': 'Advisory (10 hours/week) - Starting soon',
        }
        # Check exact match first
        if time_str_lower in time_map:
            return time_map[time_str_lower]
        # Check if it contains keywords
        if 'full_time' in time_str_lower or 'fulltime' in time_str_lower:
            if 'soon' in time_str_lower:
                return 'Full-time (40+ hours/week) - Starting soon'
            elif 'now' in time_str_lower or 'current' in time_str_lower:
                return 'Full-time (40+ hours/week) - Currently active'
            else:
                return 'Full-time (40+ hours/week)'
        elif 'part_time' in time_str_lower or 'parttime' in time_str_lower:
            if 'soon' in time_str_lower:
                return 'Part-time (20-30 hours/week) - Starting soon'
            elif 'now' in time_str_lower or 'current' in time_str_lower:
                return 'Part-time (20-30 hours/week) - Currently active'
            else:
                return 'Part-time (20-30 hours/week)'
        elif 'advisory' in time_str_lower:
            if 'soon' in time_str_lower:
                return 'Advisory (10 hours/week) - Starting soon'
            elif 'now' in time_str_lower or 'current' in time_str_lower:
                return 'Advisory (10 hours/week) - Currently active'
            else:
                return 'Advisory (10 hours/week)'
        # If no match, return as-is (might be custom text)
        return time_str
    
    founder_a_time_formatted = format_time_commitment(founder_a_time)
    founder_b_time_formatted = format_time_commitment(founder_b_time)
    
    # Get jurisdiction config for currency, governing law, etc.
    jurisdiction = vesting_terms.get('jurisdiction', 'other')
    juris_cfg = _get_jurisdiction_config(jurisdiction)
    currency_symbol = juris_cfg['currency_symbol']
    currency_code = juris_cfg['currency_code']
    
    # Format capital table (jurisdiction-aware currency)
    capital_table = f"""| Founder | Cash ({currency_code}) | Assets/IP | Loans/Guarantees | Total Value ({currency_code}) |
|---------|------------|-----------|------------------|------------------|
| {founder_a_name} | {currency_symbol}{founder_a_capital:,.0f} | [Assets/IP value] | [Loans/Guarantees] | {currency_symbol}{founder_a_capital:,.0f} |
| {founder_b_name} | {currency_symbol}{founder_b_capital:,.0f} | [Assets/IP value] | [Loans/Guarantees] | {currency_symbol}{founder_b_capital:,.0f} |"""
    
    if founder_a_capital == 0 and founder_b_capital == 0:
        capital_table = "No initial capital contributions have been made by either Founder at the time of this agreement."
    
    # Format calculation breakdown table
    calc_table = _format_calculation_breakdown(breakdown, founder_a_name, founder_b_name)
    
    # Format vesting schedule
    if not has_vesting:
        vesting_schedule = "All equity allocated to the Founders shall be **fully vested** upon the Effective Date with no vesting schedule."
    else:
        vesting_schedule = f"""All equity allocated to the Founders shall be subject to vesting over a period of **{vesting_years} years** from the Effective Date (the "**Vesting Period**").

**Cliff Vesting:**
- **No shares shall vest during the first {cliff_months} months** from the Effective Date (the "**Cliff Period**").
- At the completion of {cliff_months} months, **{int(100 / vesting_years)}%** of the total allocated shares shall vest.

**Monthly Vesting:**
- After the Cliff Period, the remaining **{100 - int(100 / vesting_years)}%** of shares shall vest in **equal monthly installments over the next {vesting_years * 12 - cliff_months} months**.
- Vesting shall occur on the same day of each month as the Effective Date."""
        
        if acceleration == 'single_trigger':
            vesting_schedule += "\n\n**Single-Trigger Acceleration:** In the event of a Change of Control (acquisition, merger, asset sale), all unvested shares shall immediately vest."
        elif acceleration == 'double_trigger':
            vesting_schedule += "\n\n**Double-Trigger Acceleration:** In the event of (a) Change of Control AND (b) involuntary termination without cause within 12 months of the Change of Control, all unvested shares shall immediately vest."
    
    # Format business stage
    stage_map = {
        'idea': 'Idea',
        'pre_seed': 'Pre-seed',
        'mvp': 'MVP Development',
        'launched': 'Launched'
    }
    business_stage_formatted = stage_map.get(business_stage, business_stage or '[Business stage to be specified]')
    
    # Build advisor equity rows if advisor is allocated equity
    # Note: Different tables have different column counts
    advisor_equity_row_2col = ""  # For 2-column tables (Stakeholder, Equity %)
    advisor_equity_row_3col = ""  # For 3-column tables (Stakeholder, Equity %, Shares)
    advisor_section = ""
    if advisor_percent > 0:
        advisor_equity_row_2col = f"| {advisor_name} (Advisor) | {advisor_percent:.2f}% |"
        advisor_equity_row_3col = f"| {advisor_name} (Advisor) | {advisor_percent:.2f}% | [Number of shares] |"
        advisor_section = f"""
### Advisor Equity

**{advisor_name}** (the "Advisor") has been granted **{advisor_percent:.2f}%** equity in the Company in consideration for advisory services to be provided.

**Advisor Vesting Schedule:**
- Vesting Period: **{advisor_vesting_years} years** from the Effective Date
- Cliff Period: **{advisor_cliff_months} months** (no vesting during this period)
- After the Cliff Period, equity vests monthly over the remaining {advisor_vesting_years * 12 - advisor_cliff_months} months

*The Advisor's equity is separate from the Founders' equity and is subject to a separate Advisor Agreement.*
"""
    
    # Build comprehensive document
    document = f"""# FOUNDERS' AGREEMENT

> ## ⚠️ DRAFT — FOR DISCUSSION & LEGAL REVIEW
>
> **This is a starting-point template** generated from the questionnaire responses provided by the Founders. It is intended to facilitate structured discussion and to serve as the basis for a formal legal agreement.
>
> **This document is NOT a legally binding contract on its own.** To be legally enforceable, this agreement must be:
> 1. **Reviewed and adapted by qualified legal counsel** in your jurisdiction ({juris_cfg['label']});
> 2. **Properly executed** by all Founders in accordance with applicable law (including, where required, on stamp paper, with witnesses, or before a notary);
> 3. **Filed or registered** with the relevant authorities where required (e.g., Companies House, Registrar of Companies, Secretary of State).
>
> Bracketed placeholders such as `[Number]`, `[Address]`, or `[X]` indicate fields that must be completed by the Founders before execution.

---

## PART I — KEY TERMS AT A GLANCE

| Item | Detail |
|------|--------|
| **Effective Date** | {date_str} |
| **Company (working name)** | {workspace_title or '[Company Name]'} |
| **Jurisdiction** | {juris_cfg['label']} |
| **Proposed Company Form** | {juris_cfg['company_form']} |
| **Business Stage** | {business_stage_formatted} |
| **Currency** | {currency_code} ({currency_symbol}) |

### Equity Allocation

| Stakeholder | Equity % |
|-------------|----------|
| **{founder_a_name}** | **{founder_a_percent:.2f}%** |
| **{founder_b_name}** | **{founder_b_percent:.2f}%** |
{f"| **{advisor_name}** (Advisor) | **{advisor_percent:.2f}%** |" + chr(10) if advisor_percent > 0 else ""}| **Total** | **100.00%** |

### Vesting at a Glance

- **Vesting period:** {vesting_years} years{' (no vesting — all equity fully vested at signing)' if not has_vesting else ''}
- **Cliff:** {cliff_months} months{'' if has_vesting else ' (n/a)'}
- **Acceleration:** {acceleration.replace('_', ' ').title() if acceleration and acceleration != 'none' else 'None'}

---

## PART II — THE AGREEMENT

**THIS FOUNDERS' AGREEMENT** ("Agreement") is made on this {day_str} day of {month_str}, {year_str} (the "**Effective Date**")

**BETWEEN:**

1. **{founder_a_name}**, residing at [Full Address], Email: [Email], Phone: [Phone] (hereinafter referred to as "**Founder A**")

AND

2. **{founder_b_name}**, residing at [Full Address], Email: [Email], Phone: [Phone] (hereinafter referred to as "**Founder B**")

(Founder A and Founder B are collectively referred to as the "**Founders**" and individually as a "**Founder**")

**IN RESPECT OF:**

**{workspace_title or '[COMPANY NAME]'}**, [a {juris_cfg['company_form']} / a company proposed to be incorporated as a {juris_cfg['company_form']}] having its [proposed] registered office at [Registered Address] (hereinafter referred to as the "**Company**").

---

## RECITALS

**WHEREAS:**

A. The Founders have agreed to jointly establish and operate the Company for the purpose of {business_description or '[brief business description]'};

B. The Company is engaged (or will engage) in the business of {business_description or '[detailed business description]'} (the "**Business**");

C. The Founders desire to set forth their respective rights, duties, obligations, and the terms governing their relationship with each other and with the Company;

D. The Founders have agreed in principle to allocate equity in the Company as follows: **{founder_a_name}** shall hold **{founder_a_percent:.2f}%** and **{founder_b_name}** shall hold **{founder_b_percent:.2f}%** of the total founding equity{f", with **{advisor_percent:.2f}%** allocated to **{advisor_name}** (Advisor)" if advisor_percent > 0 else ""}, in each case subject to the vesting schedule and other terms set forth herein;

E. This Agreement is intended to govern the relationship between the Founders pending (and to be superseded or supplemented by) a formal Shareholders' Agreement, Articles of Association/Incorporation, and any related instruments executed under {juris_cfg['governing_law']}.

**NOW, THEREFORE**, in consideration of the mutual covenants and agreements contained herein and for other good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, the Parties agree as follows:

---

## 1. PURPOSE AND BUSINESS CONCEPT

### 1.1 Business Description
The Company shall engage in {business_description or '[detailed description of business, products/services, target market, and value proposition]'}.

### 1.2 Business Stage
As of the Effective Date, the Company is at the **{business_stage_formatted}** stage.

### 1.3 Business Objectives
The Founders agree to work in good faith towards the following key objectives:
- [Objective 1 — to be specified]
- [Objective 2 — to be specified]
- [Objective 3 — to be specified]

---

## 2. EQUITY ALLOCATION AND OWNERSHIP

### 2.1 Initial Equity Distribution

The founding equity of the Company shall be allocated among the stakeholders as follows:

| Stakeholder | Equity Percentage | Number of Shares |
|-------------|-------------------|------------------|
| {founder_a_name} | {founder_a_percent:.2f}% | [Number of shares] |
| {founder_b_name} | {founder_b_percent:.2f}% | [Number of shares] |
{advisor_equity_row_3col + chr(10) if advisor_equity_row_3col else ""}| **Total** | **100%** | **[Total Founding Shares]** |

### 2.2 Basis of Equity Distribution
The equity split has been determined based on a structured assessment of the following factors:
- Time commitment and availability
- Capital contribution
- Domain expertise and technical skills
- Risk undertaken (including opportunity cost of leaving employment)
- Network and connections
- Idea origination and pre-existing intellectual property
- Roles and responsibilities

The detailed breakdown is provided in **Schedule A (Equity Calculation Matrix)**.

### 2.3 Share Capital
The {('authorized share capital' if jurisdiction in ['us', 'uk'] else 'share capital')} of the Company [is / shall be on incorporation] {currency_symbol}[Amount] divided into [Number] shares of {currency_symbol}[Face Value] each. The Founders agree to subscribe to the shares allocated to them in accordance with Section 2.1 upon incorporation or, if the Company is already incorporated, upon execution of the necessary share-issuance documents.

### 2.4 Future Dilution
The Founders acknowledge and agree that their equity percentages may be diluted in the future due to:
- Employee stock option pools (ESOP / EMI / equivalent)
- Fundraising rounds (angel, seed, venture capital)
- Strategic advisor or director grants
- Convertible instruments (SAFEs, convertible notes, CCDs, etc.)

Unless otherwise agreed in writing by all Founders, dilution shall be borne **proportionally** by all existing shareholders.

> 💡 **Note:** Specific share-issuance mechanics, classes of shares (ordinary, preferred, etc.), and pre-emption rights should be detailed in a formal Shareholders' Agreement and the Company's Articles. Founders are advised to consult counsel before issuing any shares to non-founders.

---

## 3. VESTING SCHEDULE

### 3.1 Vesting Period
{vesting_schedule}
{advisor_section}
### 3.2 Continuous Service Requirement
Vesting is contingent upon each Founder's continuous active involvement and service with the Company. Any leave of absence exceeding [30 / 60 / 90] consecutive days may pause vesting, subject to the mutual agreement of the other Founder(s).

### 3.3 Forfeiture of Unvested Shares
If a Founder ceases to be actively involved with the Company for any reason before full vesting, all unvested shares shall be forfeited and shall revert to the Company (or to a treasury / pool as agreed) on terms to be set out in a separate share buy-back or transfer instrument.

### 3.4 Tax Considerations
{('**IMPORTANT — 83(b) Election (US):** If the Company is a US entity and the Founders receive restricted stock subject to vesting, each Founder is strongly advised to file an **83(b) election** with the IRS within **30 days** of stock issuance. Failure to do so can result in significant adverse tax consequences as shares vest.' if jurisdiction == 'us' else 'Founders should consult a tax advisor in their jurisdiction regarding the tax treatment of vesting equity, including any required filings or elections.')}

---

## 4. ROLES, RESPONSIBILITIES, AND TIME COMMITMENT

### 4.1 {founder_a_name} — {founder_a_role}

**Primary Responsibilities:**
{founder_a_resp_formatted}

**Time Commitment:**
{founder_a_time_formatted}

**Key Performance Areas:**
- [KPA 1 — to be specified]
- [KPA 2 — to be specified]

### 4.2 {founder_b_name} — {founder_b_role}

**Primary Responsibilities:**
{founder_b_resp_formatted}

**Time Commitment:**
{founder_b_time_formatted}

**Key Performance Areas:**
- [KPA 1 — to be specified]
- [KPA 2 — to be specified]

### 4.3 Chief Executive Officer
[{founder_a_name} / {founder_b_name} / To be determined] shall serve as the **Chief Executive Officer (CEO)** of the Company and shall have final decision-making authority on day-to-day operational matters, subject to the Major Decisions reserved in Section 7.

### 4.4 Commitment to the Company
Each Founder agrees, during the period of their active engagement with the Company, to:
- Devote the agreed time commitment to the Company's Business in good faith;
- Not engage in any business or activity that materially conflicts with their duties to the Company without the prior written consent of the other Founder(s);
- Maintain regular communication and attend agreed meetings.

### 4.5 Modification of Roles
Roles, responsibilities, and time commitments may be modified by **mutual written consent** of all Founders and shall be documented as a written amendment to this Agreement.

---

## 5. CAPITAL CONTRIBUTIONS

### 5.1 Initial Capital
The Founders have made (or have agreed to make) the following capital contributions to the Company:

{capital_table}

### 5.2 Future Capital Requirements
Any future capital requirements shall be addressed by the Founders in good faith. Possible mechanisms include:
- Pro-rata cash contributions by the Founders;
- Founder loans on documented commercial terms (interest rate, repayment schedule);
- Raising external capital through equity or convertible instruments.

No Founder shall be obligated to contribute additional capital beyond Section 5.1 except by their own written consent.

### 5.3 Personal Guarantees
Where any Founder has provided (or proposes to provide) a personal guarantee in respect of Company obligations, the details shall be set out in **Schedule B** and the Founders shall agree on appropriate compensation, indemnity, or rebalancing.

### 5.4 Reimbursement
Founders shall be reimbursed for reasonable, documented business expenses incurred on behalf of the Company, in accordance with an expense policy adopted by the Founders.

---

## 6. INTELLECTUAL PROPERTY

> 💡 **Strongly Recommended:** A standalone **IP Assignment Agreement** (sometimes called a Proprietary Information and Inventions Agreement, or PIIA) should be executed by each Founder in favour of the Company. The clauses below set out the parties' intent, but a separate IP assignment is the standard market practice and is more readily enforceable.

### 6.1 Assignment of Company IP
All intellectual property created by any Founder in the course of, or in connection with, the Business — including but not limited to:
- ideas, inventions, and innovations;
- source code, algorithms, technical documentation, and architecture;
- designs, brand assets, trademarks, and copyrighted works;
- business processes, methodologies, and trade secrets;
- customer lists, data, and databases;

— shall be the **sole and exclusive property of the Company**, and each Founder hereby assigns (and agrees to formally assign) all right, title, and interest in such intellectual property to the Company.

### 6.2 Pre-Existing IP
{_format_ip_statement(startup_context, founder_a_name, founder_b_name)}

### 6.3 IP Assignment Documentation
Each Founder agrees to execute, and to procure the execution of, all such further deeds, assignments, and instruments as the Company may reasonably require to perfect the Company's ownership of the intellectual property described in this Section 6.

### 6.4 Third-Party IP and Open Source
Founders warrant that they have not, and will not, incorporate any third-party intellectual property (including open-source software with restrictive licensing, such as copyleft licenses) into the Company's products without proper licensing and disclosure to the other Founder(s).

---

## 7. DECISION-MAKING AND GOVERNANCE

### 7.1 Day-to-Day Decisions
Operational decisions falling within a Founder's defined role and responsibilities (Section 4) may be made by that Founder independently.

### 7.2 Major Decisions Requiring Unanimous Consent
The following decisions ("**Major Decisions**") require the prior written consent of **all Founders**:
- Any change to the equity structure or issuance of new shares;
- Any fundraising, debt financing, or sale of material assets;
- Hiring or termination of any C-level executive;
- Any material change to the business model, strategic direction, or pivot;
- Entry into any contract with a value exceeding {currency_symbol}[Amount] or a duration exceeding [12] months;
- Any sale, merger, acquisition, or change of control of the Company;
- Admission of any new founder or co-founder;
- Amendment of this Agreement, the Articles, or any Shareholders' Agreement;
- Voluntary dissolution or winding up of the Company.

### 7.3 Deadlock Resolution
In the event of a genuine deadlock on a Major Decision:
1. The Founders shall first engage in good-faith direct negotiation for not less than 15 days;
2. If unresolved, the matter shall be referred to a mutually agreed independent advisor or mediator;
3. If still unresolved, the matter shall be resolved in accordance with Section 12 (Dispute Resolution).

### 7.4 Board Composition
Upon incorporation (or if already incorporated), the Board of Directors shall initially consist of:
- {founder_a_name}
- {founder_b_name}
- [Independent Director — if applicable]

Board meeting procedures, quorum, and voting requirements shall be set out in the Articles and any Shareholders' Agreement.

---

## 8. COMPENSATION AND BENEFITS

### 8.1 Founder Compensation
Until the Company achieves [a defined revenue or funding milestone], each Founder shall:
- [Draw no salary / Draw a nominal salary of {currency_symbol}[Amount] per month].

After such milestone, salaries and benefits shall be determined by the Board (or by unanimous Founder consent) based on:
- the Company's financial position;
- market benchmarks for similar roles in the Company's jurisdiction; and
- each Founder's responsibilities and performance.

### 8.2 Reimbursements
All reasonable, documented business expenses shall be reimbursed in accordance with the Company's expense policy.

### 8.3 Benefits
Founders shall be entitled to such benefits (e.g., health insurance, professional development) as the Founders may unanimously agree from time to time.

---

## 9. CONFIDENTIALITY

### 9.1 Confidential Information
"**Confidential Information**" means any non-public information disclosed by, or learned in the course of working with, the Company or another Founder, including:
- business plans, strategies, and financial projections;
- customer, supplier, and investor lists;
- technical specifications, source code, and trade secrets;
- marketing plans, pricing, and product roadmaps;
- any information marked, or reasonably understood to be, confidential.

### 9.2 Obligations
Each Founder agrees to:
- maintain strict confidentiality during their engagement with the Company and after it ends;
- use Confidential Information solely for the purposes of the Business;
- not disclose Confidential Information to any third party without prior written consent of the other Founder(s);
- return or securely destroy all Confidential Information upon ceasing involvement with the Company.

### 9.3 Exceptions
Confidentiality obligations do not apply to information that:
- is or becomes publicly available without breach of this Agreement;
- was demonstrably known to the Founder before disclosure;
- is independently developed without reference to Confidential Information; or
- is required to be disclosed by law, court order, or competent regulatory authority (with prompt notice to the Company where lawfully permitted).

### 9.4 Duration
The obligations in this Section 9 shall survive for **[3 / 5] years** after a Founder ceases to be involved with the Company, except that obligations relating to **trade secrets** shall continue for so long as the information remains a trade secret.

---

## 10. NON-SOLICITATION AND NON-COMPETE

> ⚠️ **Jurisdiction-specific note:** {juris_cfg['non_compete_note']}

### 10.1 Non-Solicitation (Primary Restriction)
For a period of **[12] months** following the date a Founder ceases to be involved with the Company (the "**Restricted Period**"), such Founder shall not, directly or indirectly:
- solicit for employment, hire, or engage any then-current employee or contractor of the Company;
- solicit, divert, or take away any then-current customer, client, or strategic partner of the Company; or
- induce any supplier or partner to terminate or materially adversely modify its relationship with the Company.

### 10.2 Non-Compete (Reasonable Restriction)
During the Restricted Period, each Founder agrees not to, directly or indirectly, **engage in a business that is materially competitive** with the Business of the Company within [Geographic Scope]. The Founders acknowledge that this restriction is intended to be the minimum necessary to protect the Company's legitimate interests.

> Founders should review this clause with local counsel. In jurisdictions where post-employment non-compete restrictions are not enforceable (e.g., California, India), this Section 10.2 shall not apply, and the parties shall rely on Sections 6 (IP), 9 (Confidentiality), and 10.1 (Non-Solicit).

### 10.3 Reasonableness and Severability
The Founders acknowledge that the restrictions in this Section 10 are reasonable in scope, duration, and geography. If any restriction is held to be unenforceable, it shall be modified to the minimum extent necessary to make it enforceable, and the remaining restrictions shall continue in full force.

---

## 11. FOUNDER DEPARTURE AND EXIT

### 11.1 Voluntary Resignation
A Founder may voluntarily cease their active involvement with the Company by giving the other Founder(s) at least **[60] days' written notice**. On such departure:
- all unvested shares shall be forfeited (subject to Section 11.4);
- vested shares shall be subject to the buy-back rights set out in Section 11.3.

### 11.2 Termination for Cause
A Founder's involvement may be terminated by the unanimous written decision of the other Founder(s) for "**Cause**", which means:
- material and uncured breach of this Agreement;
- gross negligence, fraud, or wilful misconduct;
- conviction of an offence involving moral turpitude;
- material breach of fiduciary duties;
- prolonged unjustified absence.

On termination for Cause:
- all unvested shares are forfeited;
- the Company may exercise the buy-back rights in Section 11.3 at the **Bad Leaver** valuation.

### 11.3 Share Buy-Back
On a Founder's departure, the Company (or the remaining Founders) shall have the right (but not the obligation) to buy back the departing Founder's vested shares:

- **Valuation:** The price shall be the most recent fair value of the shares as determined by [(a) the most recent priced funding round / (b) an independent valuer / (c) a method to be agreed]. Pre-revenue, the price may be the lower of book value and fair market value.
- **Payment Terms:** Lump sum within [90] days, or in instalments over [12 / 24] months as the parties may agree.

### 11.4 Good Leaver vs. Bad Leaver
- **Good Leaver** (death, permanent disability, departure by mutual agreement): retains all vested shares; the Founders may, at their discretion, accelerate vesting of an additional [X]% of unvested shares.
- **Bad Leaver** (termination for Cause, material breach): forfeits all unvested shares; vested shares may be bought back at the lower of cost and fair value.

### 11.5 Drag-Along and Tag-Along
Drag-along and tag-along rights, anti-dilution rights, rights of first refusal, and other transfer restrictions shall be set out in a formal **Shareholders' Agreement** to be executed alongside or following this Agreement.

---

## 12. DISPUTE RESOLUTION

### 12.1 Good Faith Negotiation
The Founders shall first attempt to resolve any dispute arising out of or in connection with this Agreement through good-faith direct negotiation for a period of at least 15 days.

### 12.2 Mediation
If negotiation fails, the Founders shall attempt mediation by a mutually agreed mediator. The cost of mediation shall be shared equally.

### 12.3 Arbitration
Any dispute not resolved through mediation shall be referred to and finally resolved by arbitration:
- **Governing law of arbitration:** {juris_cfg['arbitration_act']};
- **Seat of arbitration:** {juris_cfg['arbitration_seat']};
- **Language:** English;
- **Number of arbitrators:** [1 / 3];
- **Arbitral institution / rules:** [To be specified — e.g., LCIA, ICC, AAA, MCIA, SIAC].

### 12.4 Interim Relief
Notwithstanding the agreement to arbitrate, either party may seek interim or injunctive relief from a court of competent jurisdiction in respect of breaches of confidentiality, intellectual property, or non-solicitation obligations.

---

## 13. REPRESENTATIONS AND WARRANTIES

### 13.1 Each Founder represents and warrants that:
- they have full legal capacity and authority to enter into this Agreement;
- entering into this Agreement does not breach any other contract, employment obligation, or fiduciary duty owed by them to a third party;
- they have disclosed all material information relevant to the other Founder(s) and the Company, including any prior agreements relating to the Business;
- any capital contributions are from lawful sources and made in compliance with applicable laws;
- they have not knowingly misrepresented their skills, experience, or material connections.

### 13.2 No Encumbrances
Each Founder warrants that, on issuance, their shares shall be free from any liens, charges, or encumbrances, save as may be expressly disclosed in **Schedule B**.

---

## 14. JURISDICTION-SPECIFIC PROVISIONS — {juris_cfg['label'].upper()}

### 14.1 Governing Company Law
The Company is intended to be incorporated as (or is currently) a **{juris_cfg['company_form']}** under {juris_cfg['companies_act']}.

### 14.2 Tax and Regulatory Matters
{juris_cfg['tax_note']}

### 14.3 Recommended Companion Steps
{_format_jurisdiction_recommendations(jurisdiction)}

> ⚠️ The recommendations above are **not exhaustive** and are no substitute for specific advice from local counsel and a tax advisor.

---

## 15. GENERAL PROVISIONS

### 15.1 Entire Agreement
This Agreement, together with any Schedules, constitutes the entire understanding between the Founders in respect of its subject matter and supersedes all prior discussions, communications, and agreements (whether oral or written), save for any binding pre-existing confidentiality undertakings.

### 15.2 Amendments
This Agreement may only be amended by a written instrument signed by all Founders.

### 15.3 Severability
If any provision is held invalid or unenforceable by a court or arbitrator of competent jurisdiction, the remaining provisions shall remain in full force, and the parties shall negotiate in good faith to replace the affected provision with one that achieves, to the extent permitted, the original commercial intent.

### 15.4 Waiver
A failure or delay by any party to enforce any right under this Agreement shall not constitute a waiver of such right.

### 15.5 Governing Law
This Agreement shall be governed by, and construed in accordance with, **{juris_cfg['governing_law']}**.

### 15.6 Jurisdiction
Subject to the arbitration clause in Section 12, **{juris_cfg['court_jurisdiction']}** shall have exclusive jurisdiction in respect of any matter arising out of this Agreement.

### 15.7 Notices
All notices shall be in writing and delivered to the addresses set out above (or such other address as a party may notify the others). Notices may be served by hand, recognised courier, or email (with delivery receipt).

### 15.8 Counterparts
This Agreement may be executed in any number of counterparts, including by electronic signature, each of which shall be deemed an original and which together shall constitute one and the same Agreement.

### 15.9 Successors and Assigns
This Agreement shall bind, and benefit, the parties and their respective heirs, legal representatives, and permitted assigns. No Founder may assign their rights or obligations under this Agreement without the prior written consent of the other Founder(s).

---

## SCHEDULE A — EQUITY CALCULATION MATRIX

{calc_table}

---

## SCHEDULE B — DISCLOSURES

[Disclosures regarding pre-existing IP, personal guarantees, side projects, conflicting obligations, or other material matters should be set out here.]

---

## ✍️ EXECUTION

> 💡 **Recommended:** Founders should consider executing this Agreement using a reputable e-signature platform (e.g., DocuSign, Dropbox Sign, Adobe Acrobat Sign) to maintain a tamper-evident audit trail. Where local law requires (e.g., stamp duty, notarisation, or wet-ink signatures), follow those requirements.

**IN WITNESS WHEREOF**, the Founders have executed this Agreement on the date first written above.

**FOUNDER A**

Signature: ___________________________  
Name: {founder_a_name}  
Date: ___________________________  
Place: ___________________________

Witness:  
Name: ___________________________  
Signature: ___________________________  
Address: ___________________________


**FOUNDER B**

Signature: ___________________________  
Name: {founder_b_name}  
Date: ___________________________  
Place: ___________________________

Witness:  
Name: ___________________________  
Signature: ___________________________  
Address: ___________________________

---

## ✅ "BEFORE YOU SIGN" CHECKLIST

**Both Founders should verify each of these before treating this document as final:**

- [ ] **Legal review completed** by a qualified lawyer in {juris_cfg['label']}
- [ ] All `[bracketed placeholders]` have been filled in (addresses, amounts, milestones, etc.)
- [ ] A **separate IP Assignment Agreement** has been (or will be) executed
- [ ] A formal **Shareholders' Agreement** is planned or executed (especially before raising external capital)
- [ ] The Company has been (or will be) properly incorporated and registered
- [ ] Tax implications have been reviewed with an accountant or tax advisor
{('- [ ] **83(b) election** filed within 30 days of stock issuance (US Founders ONLY — CRITICAL)' + chr(10)) if jurisdiction == 'us' else ''}{('- [ ] Stamp duty paid as per applicable State Stamp Act (India)' + chr(10)) if jurisdiction == 'india' else ''}{('- [ ] Companies House filings prepared for any share issuance (UK)' + chr(10)) if jurisdiction == 'uk' else ''}- [ ] All Founders have **read and understood** this document in full
- [ ] Each Founder has had the opportunity to seek **independent legal advice**
- [ ] All material disclosures (Schedule B) have been completed honestly

---

## 📋 IMPORTANT DISCLAIMER

**This document is a starting-point template** generated by Guild Space based on the questionnaire responses provided by the Founders.

It is intended to:
- Help Founders structure a productive conversation about equity, roles, and exit;
- Provide a reasonable starting framework for a legal agreement;
- Capture the Founders' shared understanding at the time of generation.

It is **NOT**:
- A substitute for advice from a qualified lawyer in your jurisdiction;
- A guarantee of legal enforceability — laws vary significantly by jurisdiction and over time;
- A complete legal solution — a Shareholders' Agreement, IP Assignment, and proper corporate documents are typically also required.

Guild Space and its affiliates make **no warranty** that this document is suitable for your specific circumstances, complies with the laws of your jurisdiction, or will be enforced as drafted. Use of this template is at your own risk. **Always consult qualified legal counsel before executing.**

---

**Document Version:** 2.0 (Jurisdiction-Aware)  
**Generated:** {date_str}  
**Jurisdiction:** {juris_cfg['label']}  
**Prepared Using:** Guild Space — Founders' Agreement Template
"""
    
    return document


def generate_docx(document_content: str, founder_a_name: str, founder_b_name: str) -> io.BytesIO:
    """
    Generate a DOCX file from the document content with professional formatting.
    
    Args:
        document_content: The MMD document content
        founder_a_name: Name of founder A
        founder_b_name: Name of founder B
    
    Returns:
        BytesIO object containing the DOCX file
    """
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Founders' Agreement (Draft Template)"
    doc.core_properties.author = f"{founder_a_name} & {founder_b_name}"
    
    # Set default font to Times New Roman (professional legal document font)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set heading fonts
    for heading_level in range(1, 10):
        heading_style = doc.styles[f'Heading {heading_level}']
        heading_font = heading_style.font
        heading_font.name = 'Times New Roman'
        if heading_level == 1:
            heading_font.size = Pt(18)
            heading_font.bold = True
        elif heading_level == 2:
            heading_font.size = Pt(14)
            heading_font.bold = True
        elif heading_level == 3:
            heading_font.size = Pt(12)
            heading_font.bold = True
    
    # Helper function to clean ALL HTML/Markdown tags from text for tables
    def clean_for_table(text):
        """Remove ALL HTML and markdown formatting, return plain text."""
        import re
        if not text:
            return text
        # Remove ALL HTML tags
        text = re.sub(r'<[^>]+>', '', text)
        # Remove markdown bold markers
        text = text.replace('**', '')
        return text.strip()
    
    # Helper function to format text with placeholder highlighting
    def add_text_with_placeholders(paragraph, text, is_table_cell=False):
        """Add text to paragraph, highlighting placeholders in yellow."""
        import re
        
        if is_table_cell:
            # For table cells, use completely clean plain text with proper DOCX formatting
            # First remove ALL HTML and markdown
            text = clean_for_table(text)
            # Split by placeholder pattern (matches [text] format)
            parts = re.split(r'(\[[^\]]+\])', text)
            for part in parts:
                if part.startswith('[') and part.endswith(']'):
                    run = paragraph.add_run(part)
                    run.font.highlight_color = 7  # Yellow highlight
                    run.font.color.rgb = RGBColor(128, 0, 0)  # Dark red text
                    run.font.bold = True
                elif part:
                    run = paragraph.add_run(part)
            return
        
        # For regular paragraphs (not table cells)
        # First remove any HTML tags
        text = re.sub(r'<[^>]+>', '', text)
        # Split by placeholder pattern (matches [text] format)
        parts = re.split(r'(\[[^\]]+\])', text)
        for part in parts:
            if part.startswith('[') and part.endswith(']'):
                # Highlight placeholder
                run = paragraph.add_run(part)
                run.font.highlight_color = 7  # Yellow highlight
                run.font.color.rgb = RGBColor(128, 0, 0)  # Dark red text
                run.font.bold = True
            elif part:
                # Regular text - handle markdown bold markers (**text**)
                bold_parts = part.split('**')
                for i, bold_part in enumerate(bold_parts):
                    if bold_part:
                        run = paragraph.add_run(bold_part)
                        if i % 2 == 1:  # Odd indices are bold (between ** markers)
                            run.bold = True
    
    # Parse markdown and convert to DOCX
    lines = document_content.split('\n')
    current_table_rows = []
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line - add spacing
            if current_table_rows:
                # Process accumulated table
                if len(current_table_rows) > 1:
                    table = doc.add_table(rows=len(current_table_rows), cols=len(current_table_rows[0]))
                    table.style = 'Light Grid Accent 1'
                    for i, row_data in enumerate(current_table_rows):
                        for j, cell_data in enumerate(row_data):
                            cell = table.rows[i].cells[j]
                            # Clear existing content
                            cell.text = ''
                            # Process cell content as table cell (plain text, no bold markdown)
                            paragraph = cell.paragraphs[0]
                            add_text_with_placeholders(paragraph, cell_data.strip(), is_table_cell=True)
                            # Set font for table cells
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(11)
                current_table_rows = []
            doc.add_paragraph()
            continue
        
        if line.startswith('|') and '|' in line[1:]:
            # Table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:  # Ignore separator rows (all dashes)
                if not all(c.replace('-', '').strip() == '' for c in cells):
                    current_table_rows.append(cells)
            continue
        
        # Process any accumulated table before non-table content
        if current_table_rows:
            if len(current_table_rows) > 1:
                table = doc.add_table(rows=len(current_table_rows), cols=len(current_table_rows[0]))
                table.style = 'Light Grid Accent 1'
                for i, row_data in enumerate(current_table_rows):
                    for j, cell_data in enumerate(row_data):
                        cell = table.rows[i].cells[j]
                        # Clear existing content
                        cell.text = ''
                        # Process cell content as table cell (plain text, no bold markdown)
                        paragraph = cell.paragraphs[0]
                        add_text_with_placeholders(paragraph, cell_data.strip(), is_table_cell=True)
                        # Set font for table cells
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
            current_table_rows = []
        
        if line.startswith('# '):
            # Main title
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(20)
                run.font.bold = True
        elif line.startswith('## '):
            # Section heading
            p = doc.add_heading(line[3:], level=1)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
        elif line.startswith('### '):
            # Subsection heading
            p = doc.add_heading(line[4:], level=2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
        elif line.startswith('**') and line.endswith('**'):
            # Bold paragraph
            p = doc.add_paragraph()
            add_text_with_placeholders(p, line[2:-2])
        elif line.startswith('- '):
            # Bullet point
            p = doc.add_paragraph(style='List Bullet')
            add_text_with_placeholders(p, line[2:])
        elif line.startswith('---'):
            # Horizontal rule - add a line break
            p = doc.add_paragraph('─' * 50)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)
        elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            # Italic text (disclaimer)
            p = doc.add_paragraph()
            add_text_with_placeholders(p, line[1:-1])
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(100, 100, 100)
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            add_text_with_placeholders(p, line)
    
    # Process any remaining table
    if current_table_rows:
        if len(current_table_rows) > 1:
            table = doc.add_table(rows=len(current_table_rows), cols=len(current_table_rows[0]))
            table.style = 'Light Grid Accent 1'
            for i, row_data in enumerate(current_table_rows):
                for j, cell_data in enumerate(row_data):
                    cell = table.rows[i].cells[j]
                    # Clear existing content
                    cell.text = ''
                    # Process cell content as table cell (plain text, no bold markdown)
                    paragraph = cell.paragraphs[0]
                    add_text_with_placeholders(paragraph, cell_data.strip(), is_table_cell=True)
                    # Set font for table cells
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def generate_pdf(document_content: str, founder_a_name: str, founder_b_name: str) -> io.BytesIO:
    """
    Generate a PDF file from the document content with professional formatting.
    
    Args:
        document_content: The MMD document content
        founder_a_name: Name of founder A
        founder_b_name: Name of founder B
    
    Returns:
        BytesIO object containing the PDF file
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=72)
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Professional font: Times-Roman (built-in, professional legal document font)
    # Custom styles with Times-Roman
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName='Times-Bold',
        fontSize=20,
        spaceAfter=24,
        alignment=1,  # Center
        leading=24,
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontName='Times-Bold',
        fontSize=14,
        spaceBefore=18,
        spaceAfter=10,
        leading=16,
    )
    
    subheading_style = ParagraphStyle(
        'CustomSubheading',
        parent=styles['Heading3'],
        fontName='Times-Bold',
        fontSize=12,
        spaceBefore=12,
        spaceAfter=8,
        leading=14,
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=11,
        spaceAfter=10,
        leading=13,
        leftIndent=0,
        rightIndent=0,
    )
    
    disclaimer_style = ParagraphStyle(
        'Disclaimer',
        parent=styles['Normal'],
        fontName='Times-Italic',
        fontSize=9,
        spaceAfter=8,
        leading=11,
        textColor=colors.HexColor('#666666'),
    )
    
    # Helper function to clean text for table cells (plain text, no HTML or markdown)
    def clean_for_table_pdf(text):
        """Clean text for table cells - remove all HTML and markdown, keep plain text."""
        import re
        if not text:
            return text
        # Remove ALL HTML tags
        text = re.sub(r'<[^>]+>', '', text)
        # Remove markdown bold markers but keep the text
        text = text.replace('**', '')
        return text.strip()
    
    # Helper function to convert markdown bold to HTML bold properly
    def convert_markdown_bold_to_html(text):
        """Convert **text** to <b>text</b> properly."""
        import re
        # Use regex to find pairs of ** and convert to <b>...</b>
        # Pattern: **anything** (non-greedy)
        text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
        return text
    
    # Helper function to format placeholders with highlighting (for Paragraphs)
    def format_text_with_placeholders(text, for_table=False):
        """Format text, highlighting placeholders. If for_table, return plain text."""
        import re
        
        if for_table:
            # For tables, return completely clean plain text
            return clean_for_table_pdf(text)
        
        # For paragraphs, use HTML formatting for ReportLab
        # First remove any stray HTML tags
        text = re.sub(r'<font[^>]*>', '', text)
        text = re.sub(r'</font>', '', text)
        
        # Replace placeholders with highlighted HTML for PDF
        def replace_placeholder(match):
            placeholder_text = match.group(0)
            # Use dark red text and bold for placeholders
            return f'<font color="#800000"><b>{placeholder_text}</b></font>'
        
        text = re.sub(r'\[[^\]]+\]', replace_placeholder, text)
        
        # Convert markdown bold to HTML bold properly
        text = convert_markdown_bold_to_html(text)
        
        return text
    
    # Build document
    story = []
    current_table_rows = []
    
    lines = document_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Process accumulated table before empty line
            if current_table_rows and len(current_table_rows) > 1:
                # Create table
                table_data = []
                for row in current_table_rows:
                    # Format each cell - use plain text for tables
                    formatted_row = []
                    for cell in row:
                        formatted_cell = format_text_with_placeholders(cell, for_table=True)
                        formatted_row.append(formatted_cell)
                    table_data.append(formatted_row)
                
                table = Table(table_data, colWidths=[inch * 1.5] * len(current_table_rows[0]))
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E8E8E8')),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(table)
                story.append(Spacer(1, 12))
            current_table_rows = []
            story.append(Spacer(1, 12))
            continue
        
        if line.startswith('|') and '|' in line[1:]:
            # Table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:  # Ignore separator rows (all dashes)
                if not all(c.replace('-', '').strip() == '' for c in cells):
                    current_table_rows.append(cells)
            continue
        
        # Process any accumulated table before non-table content
        if current_table_rows and len(current_table_rows) > 1:
            table_data = []
            for row in current_table_rows:
                formatted_row = []
                for cell in row:
                    formatted_cell = format_text_with_placeholders(cell, for_table=True)
                    formatted_row.append(formatted_cell)
                table_data.append(formatted_row)
            
            table = Table(table_data, colWidths=[inch * 1.5] * len(current_table_rows[0]))
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E8E8E8')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(table)
            story.append(Spacer(1, 12))
            current_table_rows = []
        
        if line.startswith('# '):
            formatted_text = format_text_with_placeholders(line[2:])
            story.append(Paragraph(formatted_text, title_style))
        elif line.startswith('## '):
            formatted_text = format_text_with_placeholders(line[3:])
            story.append(Paragraph(formatted_text, heading_style))
        elif line.startswith('### '):
            formatted_text = format_text_with_placeholders(line[4:])
            story.append(Paragraph(formatted_text, subheading_style))
        elif line.startswith('---'):
            # Horizontal rule
            story.append(Spacer(1, 10))
            t = Table([['─' * 80]], colWidths=[6.5*inch])
            t.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#999999')),
                ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
            ]))
            story.append(t)
            story.append(Spacer(1, 10))
        elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            formatted_text = format_text_with_placeholders(line[1:-1])
            story.append(Paragraph(formatted_text, disclaimer_style))
        else:
            formatted_text = format_text_with_placeholders(line)
            story.append(Paragraph(formatted_text, normal_style))
    
    # Process any remaining table
    if current_table_rows and len(current_table_rows) > 1:
        table_data = []
        for row in current_table_rows:
            formatted_row = []
            for cell in row:
                formatted_cell = format_text_with_placeholders(cell, for_table=True)
                formatted_row.append(formatted_cell)
            table_data.append(formatted_row)
        
        table = Table(table_data, colWidths=[inch * 1.5] * len(current_table_rows[0]))
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E8E8E8')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(table)
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    return buffer


def generate_and_save_document(
    clerk_user_id: str,
    workspace_id: str,
    scenario_id: Optional[str] = None
) -> Dict[str, Any]:
    """
    Generate agreement document and save to storage.
    
    Args:
        clerk_user_id: Clerk user ID
        workspace_id: Workspace ID
        scenario_id: Optional specific scenario ID (uses current if not provided)
    
    Returns:
        Dict with document info and download URLs
    """
    founder_id = _verify_workspace_access(clerk_user_id, workspace_id)
    supabase = get_supabase()
    # Use admin client for storage operations to bypass RLS (matching document_service.py pattern)
    supabase_storage = get_supabase_admin()
    if supabase_storage is None:
        log_error("WARNING: SERVICE_ROLE_KEY not set - storage operations may fail due to RLS policies")
        log_error("Please set SUPABASE_SERVICE_ROLE_KEY in your environment variables")
        supabase_storage = supabase  # Fallback to regular client
    else:
        log_info("Using admin client (service role key) for storage upload - RLS bypassed")
    
    # Get workspace title
    workspace = supabase.table('workspaces').select('title').eq('id', workspace_id).execute()
    workspace_title = workspace.data[0].get('title', 'The Company') if workspace.data else 'The Company'
    
    # Get scenario
    if scenario_id:
        scenario_query = supabase.table('equity_scenarios').select(
            '*, founder_a:founders!founder_a_id(id, name), founder_b:founders!founder_b_id(id, name)'
        ).eq('id', scenario_id).execute()
    else:
        # Get current scenario
        scenario_query = supabase.table('equity_scenarios').select(
            '*, founder_a:founders!founder_a_id(id, name), founder_b:founders!founder_b_id(id, name)'
        ).eq('workspace_id', workspace_id).eq('is_current', True).execute()
    
    if not scenario_query.data:
        raise ValueError("No approved equity scenario found. Please complete the equity questionnaire and approve a scenario first.")
    
    scenario = scenario_query.data[0]
    
    if scenario['status'] != 'approved':
        raise ValueError("Only approved scenarios can generate documents")
    
    # Get questionnaire responses
    responses = get_questionnaire_responses(clerk_user_id, workspace_id)
    
    # Extract founder names
    founder_a = scenario.get('founder_a', {})
    founder_b = scenario.get('founder_b', {})
    founder_a_name = founder_a.get('name', 'Founder A')
    founder_b_name = founder_b.get('name', 'Founder B')
    
    # Get advisor info if workspace has an advisor
    advisor = supabase.table('workspace_participants').select(
        'user:founders!user_id(id, name)'
    ).eq('workspace_id', workspace_id).eq('role', 'ADVISOR').execute()
    
    if advisor.data and advisor.data[0].get('user'):
        scenario['advisor_name'] = advisor.data[0]['user'].get('name', 'Project Advisor')
    
    # Generate MMD document
    mmd_content = generate_mmd_document(scenario, responses, workspace_title)
    
    # Generate DOCX
    docx_buffer = generate_docx(mmd_content, founder_a_name, founder_b_name)
    
    # Generate PDF
    pdf_buffer = generate_pdf(mmd_content, founder_a_name, founder_b_name)
    
    # Upload to storage
    doc_id = str(uuid.uuid4())
    docx_path = f"{workspace_id}/agreements/{doc_id}.docx"
    pdf_path = f"{workspace_id}/agreements/{doc_id}.pdf"
    
    # Upload DOCX - fail entire operation if upload fails
    try:
        log_info(f"Uploading DOCX to storage path: {docx_path}")
        storage_response = supabase_storage.storage.from_('workspace-documents').upload(
            docx_path,
            docx_buffer.getvalue(),
            file_options={
                'content-type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'upsert': False
            }
        )
        # Handle response - Supabase Python SDK returns dict or object
        if isinstance(storage_response, dict) and storage_response.get('error'):
            error_msg = storage_response['error']
            if isinstance(error_msg, dict):
                error_msg = error_msg.get('message', str(error_msg))
            raise ValueError(f"Storage upload failed: {error_msg}")
        elif hasattr(storage_response, 'error') and storage_response.error:
            error_msg = storage_response.error
            if isinstance(error_msg, dict):
                error_msg = error_msg.get('message', str(error_msg))
            raise ValueError(f"Storage upload failed: {error_msg}")
    except Exception as e:
        log_error(f"Failed to upload DOCX: {e}")
        raise ValueError(f"Failed to upload document file: {str(e)}")
    
    # Upload PDF - fail entire operation if upload fails
    try:
        log_info(f"Uploading PDF to storage path: {pdf_path}")
        storage_response = supabase_storage.storage.from_('workspace-documents').upload(
            pdf_path,
            pdf_buffer.getvalue(),
            file_options={
                'content-type': 'application/pdf',
                'upsert': False
            }
        )
        # Handle response - Supabase Python SDK returns dict or object
        if isinstance(storage_response, dict) and storage_response.get('error'):
            error_msg = storage_response['error']
            if isinstance(error_msg, dict):
                error_msg = error_msg.get('message', str(error_msg))
            raise ValueError(f"Storage upload failed: {error_msg}")
        elif hasattr(storage_response, 'error') and storage_response.error:
            error_msg = storage_response.error
            if isinstance(error_msg, dict):
                error_msg = error_msg.get('message', str(error_msg))
            raise ValueError(f"Storage upload failed: {error_msg}")
    except Exception as e:
        log_error(f"Failed to upload PDF: {e}")
        # Clean up DOCX if PDF fails
        try:
            supabase_storage.storage.from_('workspace-documents').remove([docx_path])
        except Exception:
            pass
        raise ValueError(f"Failed to upload document file: {str(e)}")
    
    # Generate signed URLs
    docx_url = None
    pdf_url = None
    
    if docx_path:
        try:
            docx_signed = supabase_storage.storage.from_('workspace-documents').create_signed_url(docx_path, 3600)
            docx_url = docx_signed.get('signedUrl') if isinstance(docx_signed, dict) else getattr(docx_signed, 'signedUrl', None)
        except Exception as e:
            log_error(f"Failed to create signed URL for DOCX: {e}")
    
    if pdf_path:
        try:
            pdf_signed = supabase_storage.storage.from_('workspace-documents').create_signed_url(pdf_path, 3600)
            pdf_url = pdf_signed.get('signedUrl') if isinstance(pdf_signed, dict) else getattr(pdf_signed, 'signedUrl', None)
        except Exception as e:
            log_error(f"Failed to create signed URL for PDF: {e}")
    
    # Save document record
    doc_record = {
        'workspace_id': workspace_id,
        'scenario_id': scenario['id'],
        'document_content': mmd_content,
        'pdf_url': pdf_path,
        'docx_url': docx_path,
        'generated_by': founder_id,
    }
    
    # Use admin client for database insert to bypass RLS (matching document_service.py pattern)
    db_client = get_supabase_admin() or supabase
    result = db_client.table('generated_equity_documents').insert(doc_record).execute()
    
    if not result.data:
        raise ValueError("Failed to save document record")
    
    _log_audit(
        workspace_id, founder_id,
        'generate_equity_document',
        'generated_equity_document',
        result.data[0]['id']
    )
    
    return {
        'id': result.data[0]['id'],
        'mmd_content': mmd_content,
        'docx_url': docx_url,
        'pdf_url': pdf_url,
        'generated_at': result.data[0]['generated_at'],
    }


def get_document(
    clerk_user_id: str,
    workspace_id: str,
    document_id: str
) -> Dict[str, Any]:
    """
    Get a previously generated document with fresh signed URLs.
    
    Args:
        clerk_user_id: Clerk user ID
        workspace_id: Workspace ID
        document_id: Document ID
    
    Returns:
        Document info with signed URLs
    """
    _verify_workspace_access(clerk_user_id, workspace_id)
    supabase = get_supabase()
    # Use admin client for storage operations to bypass RLS
    supabase_admin = get_supabase_admin()
    
    doc = supabase.table('generated_equity_documents').select('*').eq(
        'id', document_id
    ).eq('workspace_id', workspace_id).execute()
    
    log_info(f"get_document: fetched doc data: {doc.data}")
    
    if not doc.data:
        raise ValueError("Document not found")
    
    doc_data = doc.data[0]
    
    # Generate fresh signed URLs using admin client
    docx_url = None
    pdf_url = None
    
    if doc_data.get('docx_url'):
        try:
            log_info(f"Generating signed URL for DOCX path: {doc_data['docx_url']}")
            docx_signed = supabase_admin.storage.from_('workspace-documents').create_signed_url(
                doc_data['docx_url'], 3600
            )
            log_info(f"DOCX signed URL response: {docx_signed}")
            docx_url = docx_signed.get('signedUrl') if isinstance(docx_signed, dict) else getattr(docx_signed, 'signedUrl', None)
        except Exception as e:
            log_error(f"Failed to create signed URL for DOCX: {e}")
    
    if doc_data.get('pdf_url'):
        try:
            log_info(f"Generating signed URL for PDF path: {doc_data['pdf_url']}")
            pdf_signed = supabase_admin.storage.from_('workspace-documents').create_signed_url(
                doc_data['pdf_url'], 3600
            )
            log_info(f"PDF signed URL response: {pdf_signed}")
            pdf_url = pdf_signed.get('signedUrl') if isinstance(pdf_signed, dict) else getattr(pdf_signed, 'signedUrl', None)
        except Exception as e:
            log_error(f"Failed to create signed URL for PDF: {e}")
    
    return {
        'id': doc_data['id'],
        'scenario_id': doc_data['scenario_id'],
        'mmd_content': doc_data['document_content'],
        'docx_url': docx_url,
        'pdf_url': pdf_url,
        'generated_at': doc_data['generated_at'],
    }


def list_documents(
    clerk_user_id: str,
    workspace_id: str
) -> list:
    """
    List all generated documents for a workspace.
    
    Args:
        clerk_user_id: Clerk user ID
        workspace_id: Workspace ID
    
    Returns:
        List of document records
    """
    _verify_workspace_access(clerk_user_id, workspace_id)
    supabase = get_supabase()
    
    docs = supabase.table('generated_equity_documents').select(
        'id, scenario_id, generated_at, version'
    ).eq('workspace_id', workspace_id).order('generated_at', desc=True).execute()
    
    return docs.data or []


def download_document(
    clerk_user_id: str,
    workspace_id: str,
    document_id: str,
    file_type: str = 'pdf'
) -> Tuple[bytes, str, str]:
    """
    Download document file content directly (proxy download).
    This avoids exposing Supabase signed URLs to the client.
    
    Args:
        clerk_user_id: Clerk user ID
        workspace_id: Workspace ID
        document_id: Document ID
        file_type: 'pdf' or 'docx'
    
    Returns:
        Tuple of (file_content, content_type, filename)
    """
    _verify_workspace_access(clerk_user_id, workspace_id)
    supabase = get_supabase()
    supabase_admin = get_supabase_admin()
    
    # Get document record
    doc = supabase.table('generated_equity_documents').select('*').eq(
        'id', document_id
    ).eq('workspace_id', workspace_id).execute()
    
    if not doc.data:
        raise ValueError("Document not found")
    
    doc_data = doc.data[0]
    
    # Determine which file to download
    if file_type == 'pdf':
        file_path = doc_data.get('pdf_url')
        content_type = 'application/pdf'
        filename = f"co-founder-agreement-{document_id[:8]}.pdf"
    elif file_type == 'docx':
        file_path = doc_data.get('docx_url')
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        filename = f"co-founder-agreement-{document_id[:8]}.docx"
    else:
        raise ValueError("Invalid file type. Use 'pdf' or 'docx'")
    
    if not file_path:
        raise ValueError(f"No {file_type.upper()} file available for this document")
    
    # Download file from Supabase storage
    try:
        log_info(f"Downloading file from storage: {file_path}")
        file_response = supabase_admin.storage.from_('workspace-documents').download(file_path)
        
        if file_response is None:
            raise ValueError("Failed to download file from storage")
        
        # file_response is bytes
        return file_response, content_type, filename
        
    except Exception as e:
        log_error(f"Failed to download file: {e}")
        raise ValueError(f"Failed to download file: {str(e)}")
